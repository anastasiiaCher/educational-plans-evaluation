import pandas as pd
import glob
import os
import docx
import re
import warnings
from tqdm import tqdm
import nltk
import pymorphy2
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import re
import gensim
from gensim.models import word2vec
from gensim.models.phrases import Phrases, Phraser
from deeppavlov import build_model, configs
from tabulate import tabulate

warnings.filterwarnings("ignore")

# nltk.download('stopwords')
# nltk.download('punkt')

# syntax_model = build_model(configs.syntax.syntax_ru_syntagrus_bert, download=True)


def find_docx(path):
    os.chdir(path)
    files = [file for file in glob.glob("*.docx")]
    return files


# достаем раздел с содержанием из РПД (из docx)
def extract_contents(files):
    data = []
    for num in tqdm(range(len(files))):
        try:
            rpd_data = []
            doc = docx.Document("RPD/%s" % files[num])
            title = [par.text for par in doc.paragraphs if par.text][4]
            for t in range(len(doc.tables)):
                table = doc.tables[t]
                keys = None
                for i, row in enumerate(table.rows):
                    if row.cells:
                        text = (cell.text for cell in row.cells if cell)
                        if i == 0:
                            keys = tuple(re.sub(r"\s+", " ", key.strip()) for key in tuple(text))
                            continue
                        if 'Наименование раздела дисциплины' in keys and 'Содержание' in keys:
                            row_data = dict(zip(keys, text))
                            rpd_data.append(row_data)
            if rpd_data:
                r_data = pd.DataFrame(rpd_data)
                r_data["raw_text"] = r_data.apply(lambda r: str(r["Наименование раздела дисциплины"]) + ". " + str(r["Содержание"]), axis=1)
                data.append({"file": files[num],
                             "title": title,
                             "content": " ".join(r_data.raw_text.tolist())})
            else:
                rpd_data = []
                for t in range(len(doc.tables)):
                    table = doc.tables[t]
                    keys = None
                    for i, row in enumerate(table.rows):
                        if row.cells:
                            text = (cell.text for cell in row.cells if cell)
                            if i == 0:
                                keys = tuple(re.sub(r"\s+", " ", key.strip()) for key in tuple(text))
                                continue
                            if 'Наименование раздела дисциплины' in keys:
                                row_data = dict(zip(keys, text))
                                rpd_data.append(row_data)
                if rpd_data:
                    r_data = pd.DataFrame(rpd_data)
                    data.append({"file": files[num],
                                 "title": title,
                                 "content": " ".join(r_data["Наименование раздела дисциплины"].tolist())})
        except ValueError:
            print(files[num])
            continue
    df = pd.DataFrame(data)
    return df


def clean_title(text):
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-zA-Zа-яА-ЯёЁ\- ]", "", text)
    text = " ".join(text.split() if not text.split()[0].isupper() and len(text.split()[0]) > 2 else text.split()[1:])
    return text.strip() if not text.isupper() else ""


def clean_text(text):
    text = re.sub("итого", "", text, flags=re.I)
    text = re.sub(r"наименование\s+раздела\s+дисциплины", "", text, flags=re.I)
    return text.strip()


def rpd_to_wordlist(text, morph_model, remove_stopwords=False, morph=True):
    # оставляем только буквенные и числовые символ
    text = re.sub("[^a-zA-Zа-яА-ЯёЁ\-]", " ", text)
    # приводим к нижнему регистру и разбиваем на слова по символу пробела
    words = text.lower().split()
    norm_words = []
    if words:
        if morph:
            norm_words = [parsed.split()[2] for parsed in morph_model(words)]
            # norm_words = [analyzer.parse(word)[0].normal_form for word in words]
            if remove_stopwords:
                # убираем стоп-слова
                stops = stopwords.words("russian")
                norm_words = [w for w in norm_words if w not in stops]
        elif remove_stopwords:
            # убираем стоп-слова
            stops = stopwords.words("russian")
            norm_words = [w for w in words if w not in stops]
        return norm_words if norm_words else words
    return words


def rpd_to_sentences(review, tokenizer, morph_model, remove_stopwords=False, morph=True):
    raw_sentences = tokenizer.tokenize(review.strip())
    sentences = []
    for raw_sentence in raw_sentences:
        if len(raw_sentence) > 0:
            sentences.append(rpd_to_wordlist(raw_sentence, morph_model, remove_stopwords, morph))
    return sentences


def tokenize_data(data):
    morph_model = build_model(configs.morpho_tagger.BERT.morpho_ru_syntagrus_bert)
    print("ok1")
    sentences = []
    tokenized_cont = []
    tokenizer = nltk.data.load('tokenizers/punkt/russian.pickle')
    for rpd in tqdm(data):
        cont = rpd_to_sentences(rpd, tokenizer, morph_model, remove_stopwords=True)
        sentences.extend(cont)
        tokenized_cont.append(cont)
    return sentences, tokenized_cont


def train_bt_model(sents):
    print("Training model...")
    # обучаем биграммную модель
    bigram = Phrases(sentences, min_count=1, threshold=5)
    bigram_phraser = Phraser(bigram)
    # обучаем триграммную модель
    trigram = Phrases(bigram[sentences], min_count=1)
    trigram_phraser = Phraser(trigram)
    return bigram_phraser, trigram_phraser


def get_good_bi_tri(sents, bigram_phraser, trigram_phraser):
    morph_model = build_model(configs.morpho_tagger.BERT.morpho_ru_syntagrus_bert)
    print("ok2")
    # синтаксические правила
    rules = {1: {"POS": {"NOUN", "X"}},
             2: {1: {"POS1": {"ADJ", "PRT"}, "POS2": {"NOUN", "X", "PROPN"}},
                 2: {"POS1": {"NOUN", "X", "PROPN"}, "POS2": {"NOUN", "X", "PROPN"}}},
             3: {1: {"POS1": {"NOUN", "X", "PROPN"}, "POS2": {"NOUN", "X", "PROPN"}, "POS3": {"NOUN", "X", "PROPN"}},
                 2: {"POS1": {"NOUN", "X", "PROPN"}, "POS2": {"ADJ", "PRT"}, "POS3": {"NOUN", "X", "PROPN"}},
                 3: {"POS1": {"ADJ", "PRT"}, "POS2": {"NOUN", "X", "PROPN"}, "POS3": {"NOUN", "X", "PROPN"}},
                 4: {"POS1": {"ADJ", "PRT"}, "POS2": {"ADJ", "PRT"}, "POS3": {"NOUN", "X", "PROPN"}}}}
    bi_and_tri = []
    for sent in sents:
        # print("\t", "sentence:", sent)
        bigrams_ = [b for b in bigram_phraser[sent] if b.count("_") == 1]
        trigrams_ = [t for t in trigram_phraser[bigram_phraser[sent]] if t.count("_") == 2]
        if bigrams_:
            bigrams_ = [phrase.split("_") for phrase in bigrams_]
            norm_bi = [[(parsed.split()[2], parsed.split()[3]) for parsed in morph_model(phrase)] for phrase in
                       bigrams_]
            norm_bi = [[t[0] for t in tt] for tt in norm_bi if
                       tt[0][1] in rules[2][2]["POS1"] and tt[1][1] in rules[2][2]["POS2"] or tt[0][1] in rules[2][1][
                           "POS1"] and tt[1][1] in rules[2][1]["POS2"]]
            # print("bigrams:", bigrams_, norm_bi, sep="\n")
            bi_and_tri.extend(norm_bi)
        if trigrams_:
            trigrams_ = [phrase.split("_") for phrase in trigrams_]
            norm_tri = [[(parsed.split()[2], parsed.split()[3]) for parsed in morph_model(phrase)] for phrase in
                        trigrams_]
            norm_tri = [[t[0] for t in tt] for tt in norm_tri if
                        tt[0][1] in rules[3][1]["POS1"] and tt[1][1] in rules[3][1]["POS2"] and tt[2][1] in rules[3][1][
                            "POS3"] or tt[0][1] in rules[3][2]["POS1"] and tt[1][1] in rules[3][2]["POS2"] and tt[2][1] in rules[3][2][
                            "POS3"] or tt[0][1] in rules[3][3]["POS1"] and tt[1][1] in rules[3][3]["POS2"] and tt[2][1] in rules[3][3][
                            "POS3"] or tt[0][1] in rules[3][4]["POS1"] and tt[1][1] in rules[3][4]["POS2"] and tt[2][1] in rules[3][4][
                            "POS3"]]
            # print("trigrams:", trigrams_, norm_tri, sep="\n")
            bi_and_tri.extend(norm_tri)
    return bi_and_tri


if __name__ == "__main__":
    # rpds = find_docx("RPD")
    # os.chdir("/home/siia/PycharmProjects/individual-tracks")
    # rpd_cont = extract_contents(rpds)
    # rpd_cont.to_excel("rpd_content.xlsx", index=False)
    # rpd_cont.title = rpd_cont.title.apply(clean_title)
    # rpd_cont["text"] = rpd_cont.apply(lambda i: "{}. {}".format(i.title, i.content), axis=1)
    # rpd_cont.text = rpd_cont.text.apply(clean_text)
    # sents, tok_cont = tokenize_data(rpd_cont.text.tolist())
    # rpd_cont["tokens"] = tok_cont
    # with open("sentences10062021.txt", 'w') as f:
    #    for s in sents:
    #        f.write(' '.join(s))
    #        f.write('\n')
    rpd_cont = pd.read_excel("rpd_content.xlsx", engine='openpyxl')
    sentences = []
    with open("sentences10062021.txt", 'r') as f:
        for line in f:
            sentences.append(line.split())
    phraser2, phraser3 = train_bt_model(sentences)
    # morph = build_model(configs.morpho_tagger.BERT.morpho_ru_syntagrus_bert)
    print(get_good_bi_tri(tokenize_data([rpd_cont.text.iloc[15]])[0], phraser2, phraser3))
    print(rpd_cont.text.iloc[15])

